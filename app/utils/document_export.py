from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Optional


def create_docx(output_path: str, arabic_text: str, translated_text: Optional[str] = None, transliterated_text: Optional[str] = None) -> None:
    """Create a DOCX file with Arabic text, translation, and transliteration"""
    # Create a new Document
    doc = Document()
    
    # Set up document styles
    styles = doc.styles
    
    # Create a style for Arabic text
    if 'Arabic' not in styles:
        arabic_style = styles.add_style('Arabic', WD_STYLE_TYPE.PARAGRAPH)
        font = arabic_style.font
        font.name = 'Arial'
        font.size = Pt(14)
        font.bold = True
        arabic_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Create a style for English text
    if 'English' not in styles:
        english_style = styles.add_style('English', WD_STYLE_TYPE.PARAGRAPH)
        font = english_style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        english_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Create a style for transliteration
    if 'Transliteration' not in styles:
        transliteration_style = styles.add_style('Transliteration', WD_STYLE_TYPE.PARAGRAPH)
        font = transliteration_style.font
        font.name = 'Courier New'
        font.size = Pt(12)
        font.italic = True
        transliteration_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Create a style for section headings
    if 'SectionHeading' not in styles:
        heading_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        font = heading_style.font
        font.name = 'Arial'
        font.size = Pt(16)
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 128)  # Navy blue
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
    
    # Add document title
    title = doc.add_paragraph("Arabic OCR Results", style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Arabic text section
    doc.add_paragraph("Original Arabic Text", style='SectionHeading')
    
    # Split Arabic text by lines and add each line
    for line in arabic_text.split('\n'):
        if line.strip():
            doc.add_paragraph(line, style='Arabic')
        else:
            doc.add_paragraph()
    
    # Add translation section if provided
    if translated_text:
        doc.add_paragraph("English Translation", style='SectionHeading')
        
        # Split translated text by lines and add each line
        for line in translated_text.split('\n'):
            if line.strip():
                doc.add_paragraph(line, style='English')
            else:
                doc.add_paragraph()
    
    # Add transliteration section if provided
    if transliterated_text:
        doc.add_paragraph("Latin Transliteration", style='SectionHeading')
        
        # Split transliterated text by lines and add each line
        for line in transliterated_text.split('\n'):
            if line.strip():
                doc.add_paragraph(line, style='Transliteration')
            else:
                doc.add_paragraph()
    
    # Save the document
    doc.save(output_path)